"""
Extraction Agent
Extracts structured data from PDF, Word, and plain-text files.
Returns a pandas DataFrame.
"""
import io
import re
import pandas as pd


class ExtractionAgent:
    """Extracts tabular data from unstructured document formats."""

    name = "ExtractionAgent"

    def run(self, state: dict) -> dict:
        file_type = state.get("file_type")
        uploaded_file = state.get("uploaded_file")
        state["agent_log"] = state.get("agent_log", [])

        if file_type not in ("pdf", "docx", "txt"):
            # Not needed for this file type
            return state

        try:
            if file_type == "pdf":
                df = self._extract_pdf(uploaded_file)
            elif file_type == "docx":
                df = self._extract_docx(uploaded_file)
            else:
                df = self._extract_txt(uploaded_file)

            state["extracted_df"] = df
            state["agent_log"].append(
                f"✅ ExtractionAgent: Extracted {df.shape[0]} rows × {df.shape[1]} columns "
                f"from {file_type.upper()} file."
            )
        except Exception as e:
            state["error"] = f"ExtractionAgent failed: {e}"

        return state

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------
    def _extract_pdf(self, uploaded_file) -> pd.DataFrame:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF extraction. Run: pip install pdfplumber")

        uploaded_file.seek(0)
        with pdfplumber.open(uploaded_file) as pdf:
            all_tables = []
            all_text_rows = []

            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        all_tables.extend(table)
                else:
                    # Fallback: extract raw text lines
                    text = page.extract_text() or ""
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            all_text_rows.append(line)

            if all_tables:
                headers = all_tables[0]
                rows = all_tables[1:]
                df = pd.DataFrame(rows, columns=headers)
            else:
                # Build DataFrame from text lines using delimiter detection
                df = self._parse_text_lines(all_text_rows)

        return df

    # ------------------------------------------------------------------
    # Word
    # ------------------------------------------------------------------
    def _extract_docx(self, uploaded_file) -> pd.DataFrame:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Run: pip install python-docx")

        uploaded_file.seek(0)
        doc = Document(uploaded_file)

        # Try tables first
        if doc.tables:
            table = doc.tables[0]
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            if data:
                df = pd.DataFrame(data[1:], columns=data[0])
                return df

        # Fallback: paragraph text
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return self._parse_text_lines(lines)

    # ------------------------------------------------------------------
    # Plain text
    # ------------------------------------------------------------------
    def _extract_txt(self, uploaded_file) -> pd.DataFrame:
        uploaded_file.seek(0)
        content = uploaded_file.read().decode("utf-8", errors="replace")
        lines = [l.strip() for l in content.splitlines() if l.strip()]
        return self._parse_text_lines(lines)

    # ------------------------------------------------------------------
    # Helper: detect delimiter and parse lines into DataFrame
    # ------------------------------------------------------------------
    def _parse_text_lines(self, lines: list) -> pd.DataFrame:
        if not lines:
            return pd.DataFrame()

        # Detect delimiter
        sample = "\n".join(lines[:20])
        delimiter = ","
        for d in [",", "\t", ";", "|"]:
            if lines[0].count(d) >= 1:
                delimiter = d
                break

        from io import StringIO
        try:
            df = pd.read_csv(StringIO("\n".join(lines)), sep=delimiter, engine="python")
            if df.shape[1] > 1:
                return df
        except Exception:
            pass

        # Last resort: single column
        return pd.DataFrame({"text": lines})
